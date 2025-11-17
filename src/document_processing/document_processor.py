"""
Document processor for parsing and chunking various document formats.

Supports: PDF, DOCX, TXT, CSV, XLS, XLSX, and other text-based formats.
"""

import hashlib
import logging
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from docx import Document as DocxDocument
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import TextNode
from pypdf import PdfReader

from src.config import settings

logger = logging.getLogger(__name__)

# Lazy import to avoid circular dependencies and API key errors during initialization
_enhancement_processor = None


def _get_enhancement_processor():
    """
    Get or create EnhancementProcessor instance (lazy initialization).

    Returns:
        EnhancementProcessor instance or None if enhancement is disabled.
    """
    global _enhancement_processor

    if not settings.doc_enhancement_enabled:
        return None

    if _enhancement_processor is None:
        try:
            from src.document_processing.enhancement_processor import (
                EnhancementProcessor,
            )

            model = settings.doc_enhancement_model or settings.openrouter_model
            _enhancement_processor = EnhancementProcessor(
                llm_model=model,
                max_tokens=settings.doc_enhancement_max_tokens,
            )
            logger.info("EnhancementProcessor initialized for document enhancement")
        except Exception as e:
            logger.warning(f"Failed to initialize EnhancementProcessor: {e}")
            return None

    return _enhancement_processor


class DocumentProcessor:
    """
    Processes documents from various formats into chunked text nodes.

    Handles PDF, DOCX, TXT, CSV, XLS, XLSX and other common document formats.
    """

    def __init__(
        self,
        chunk_size: int = settings.chunk_size,
        chunk_overlap: int = settings.chunk_overlap,
    ):
        """
        Initialize the document processor.

        Args:
            chunk_size: Size of each text chunk in characters.
            chunk_overlap: Overlap between consecutive chunks.
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Convert characters to tokens (roughly 1 token = 4 characters)
        # SentenceSplitter uses tokens, not characters
        chunk_size_tokens = chunk_size // 4
        chunk_overlap_tokens = chunk_overlap // 4

        self.splitter = SentenceSplitter(
            chunk_size=chunk_size_tokens,
            chunk_overlap=chunk_overlap_tokens,
        )

    def compute_file_hash(self, file_path: Path) -> str:
        """
        Compute SHA-256 hash of a file for deduplication.

        Args:
            file_path: Path to the file.

        Returns:
            Hex string of the file hash.
        """
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read in chunks to handle large files
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file.

        Returns:
            Extracted text content.

        Raises:
            Exception: If PDF reading fails.
        """
        try:
            reader = PdfReader(file_path)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file.

        Returns:
            Extracted text content.

        Raises:
            Exception: If DOCX reading fails.
        """
        try:
            doc = DocxDocument(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

    def extract_text_from_txt(self, file_path: Path) -> str:
        """
        Extract text from plain text file.

        Args:
            file_path: Path to text file.

        Returns:
            File content as string.

        Raises:
            Exception: If file reading fails.
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            raise

    def extract_text_from_csv(self, file_path: Path) -> str:
        """
        Extract text from CSV file.

        Args:
            file_path: Path to CSV file.

        Returns:
            CSV content formatted as text optimized for semantic search.

        Raises:
            Exception: If CSV reading fails.
        """
        try:
            # Try reading with headers first
            df = pd.read_csv(file_path)

            # Check if the file likely has no real headers (too many unnamed columns)
            unnamed_cols = [col for col in df.columns if "Unnamed" in str(col)]
            if len(unnamed_cols) > len(df.columns) / 2:
                # Re-read without treating first row as header
                df = pd.read_csv(file_path, header=None)
                has_headers = False
            else:
                has_headers = True

            # Convert DataFrame to a semantic-search-friendly format
            lines = []

            # Add filename as context
            lines.append(f"Document: {file_path.name}")
            lines.append("")

            if has_headers:
                # Format with column names for each row
                for idx, row in df.iterrows():
                    # Skip completely empty rows
                    if row.isna().all():
                        lines.append("")
                        continue

                    row_lines = []
                    for col_name, value in row.items():
                        if pd.notna(value) and str(value).strip():
                            row_lines.append(f"{col_name}: {value}")

                    if row_lines:
                        lines.extend(row_lines)
                        lines.append("")  # Empty line between rows
            else:
                # Format as simple rows (for CSVs without headers)
                for idx, row in df.iterrows():
                    # Skip rows with only empty/zero values
                    row_values = [
                        str(val)
                        for val in row
                        if pd.notna(val) and str(val).strip() and str(val) != "0"
                    ]

                    if not row_values:
                        lines.append("")  # Preserve spacing for term breaks
                        continue

                    # Check if first column looks like a section header (contains month/term)
                    first_val = str(row[0]).strip() if pd.notna(row[0]) else ""

                    if len(row_values) == 1 or "term" in first_val.lower():
                        # Section header
                        lines.append(first_val)
                        lines.append("")
                    else:
                        # Regular data row - format as "Date: Event [Notes]"
                        date_col = first_val
                        event_col = (
                            str(row[1]).strip()
                            if len(row) > 1 and pd.notna(row[1])
                            else ""
                        )
                        notes_col = (
                            str(row[2]).strip()
                            if len(row) > 2 and pd.notna(row[2])
                            else ""
                        )

                        if event_col:
                            line = f"Date: {date_col} - Event: {event_col}"
                            if notes_col:
                                line += f" - Notes: {notes_col}"
                            lines.append(line)

            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {e}")
            raise

    def extract_text_from_excel(self, file_path: Path) -> str:
        """
        Extract text from Excel file (XLS/XLSX).

        Args:
            file_path: Path to Excel file.

        Returns:
            Excel content formatted as text optimized for semantic search.

        Raises:
            Exception: If Excel reading fails.
        """
        try:
            import openpyxl

            # Load workbook (data_only=True to get calculated values, not formulas)
            wb = openpyxl.load_workbook(file_path, data_only=True)

            lines = []

            # Add filename as context
            lines.append(f"Document: {file_path.name}")
            lines.append("")

            # Process each sheet
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]

                # Skip empty sheets
                if sheet.max_row == 0 or sheet.max_column == 0:
                    continue

                # Add sheet name as section header
                lines.append(f"Sheet: {sheet_name}")
                lines.append("")

                # Read all rows into a list
                rows_data = list(sheet.iter_rows(values_only=True))

                if not rows_data:
                    continue

                # Detect if first row is likely headers
                first_row = rows_data[0]
                # Headers are likely if first row contains mostly strings and no None
                non_none_first = [v for v in first_row if v is not None]
                has_headers = (
                    len(non_none_first) > 0
                    and all(isinstance(v, str) for v in non_none_first)
                    and len(non_none_first) >= len(first_row) / 2
                )

                if has_headers:
                    # Use first row as column names
                    headers = [
                        str(h) if h is not None else f"Column{i}"
                        for i, h in enumerate(first_row)
                    ]
                    data_rows = rows_data[1:]

                    # Format each data row with column names
                    for row in data_rows:
                        # Skip completely empty rows
                        if all(v is None or str(v).strip() == "" for v in row):
                            lines.append("")
                            continue

                        row_lines = []
                        for col_name, value in zip(headers, row):
                            if value is not None and str(value).strip():
                                row_lines.append(f"{col_name}: {value}")

                        if row_lines:
                            lines.extend(row_lines)
                            lines.append("")  # Empty line between rows
                else:
                    # No headers - format as simple rows
                    for row in rows_data:
                        # Skip rows with only empty/None values
                        row_values = [
                            str(val)
                            for val in row
                            if val is not None and str(val).strip()
                        ]

                        if not row_values:
                            lines.append("")  # Preserve spacing
                            continue

                        # Check if first column looks like a section header
                        first_val = str(row[0]).strip() if row[0] is not None else ""

                        if len(row_values) == 1:
                            # Single value - likely a section header
                            lines.append(first_val)
                            lines.append("")
                        else:
                            # Multiple values - format as structured data
                            line_parts = []
                            for i, val in enumerate(row):
                                if val is not None and str(val).strip():
                                    line_parts.append(f"Col{i+1}: {val}")
                            if line_parts:
                                lines.append(" | ".join(line_parts))

                # Add spacing between sheets
                lines.append("")

            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Error reading Excel file {file_path}: {e}")
            raise

    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from file based on extension.

        Args:
            file_path: Path to the file.

        Returns:
            Extracted text content.

        Raises:
            ValueError: If file format is not supported.
            Exception: If extraction fails.
        """
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif suffix == ".docx":
            return self.extract_text_from_docx(file_path)
        elif suffix == ".txt":
            return self.extract_text_from_txt(file_path)
        elif suffix == ".csv":
            return self.extract_text_from_csv(file_path)
        elif suffix in (".xls", ".xlsx"):
            return self.extract_text_from_excel(file_path)
        else:
            # Try to read as text file
            try:
                return self.extract_text_from_txt(file_path)
            except Exception:
                raise ValueError(f"Unsupported file format: {suffix}")

    def process_document(
        self,
        file_path: Path,
        source_type: str = "manual",
        extra_metadata: Optional[Dict] = None,
    ) -> List[TextNode]:
        """
        Process a document file into chunked text nodes.

        Args:
            file_path: Path to the document file.
            source_type: Source of the document ('manual' or 'email').
            extra_metadata: Additional metadata to attach to nodes.

        Returns:
            List of TextNode objects with text chunks and metadata.

        Raises:
            Exception: If document processing fails.
        """
        logger.info(f"Processing document: {file_path}")

        try:
            # Extract text from file
            text = self.extract_text(file_path)

            if not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return []

            # Enhance document if it's structured data (CSV/Excel)
            file_type = file_path.suffix.lower()
            enhanced = False
            enhancement_count = 0

            enhancer = _get_enhancement_processor()
            if enhancer and enhancer.should_enhance(file_type):
                try:
                    logger.info(f"Enhancing {file_type} document: {file_path.name}")

                    # Parse enhancement types from settings
                    enhancement_types = [
                        t.strip()
                        for t in settings.doc_enhancement_types.split(",")
                        if t.strip()
                    ]

                    # Generate enhanced content
                    enhancement_result = enhancer.enhance_document(
                        text, file_type, enhancement_types
                    )

                    # Append enhanced content to original text
                    if enhancement_result["enhanced_text"]:
                        text = f"{text}\n\n{enhancement_result['enhanced_text']}"
                        enhanced = True
                        enhancement_count = enhancement_result["enhancement_count"]
                        logger.info(
                            f"Successfully enhanced document with {enhancement_count} enhancements"
                        )

                except Exception as e:
                    logger.error(f"Document enhancement failed: {e}")
                    # Continue with unenhanced document
                    pass

            # Compute file hash for deduplication
            file_hash = self.compute_file_hash(file_path)

            # Get file modification timestamp for versioning
            file_stat = file_path.stat()
            file_mtime = file_stat.st_mtime
            file_size = file_stat.st_size

            # Create base metadata
            metadata = {
                "filename": file_path.name,
                "file_path": str(file_path),
                "file_hash": file_hash,
                "file_mtime": file_mtime,  # Modification timestamp for versioning
                "file_size": file_size,
                "source_type": source_type,
                "file_type": file_type,
                "enhanced": enhanced,
                "enhancement_count": enhancement_count,
            }

            # Add extra metadata if provided
            if extra_metadata:
                metadata.update(extra_metadata)

            # Create Document object for LlamaIndex
            document = Document(text=text, metadata=metadata)

            # Split into chunks
            nodes = self.splitter.get_nodes_from_documents([document])

            logger.info(
                f"Successfully processed {file_path.name}: {len(nodes)} chunks created"
            )

            return nodes

        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {e}")
            raise

    def process_url(
        self,
        url: str,
        crawl_depth: int = 1,
        extra_metadata: Optional[Dict] = None,
    ) -> List[TextNode]:
        """
        Process web content from URL(s) into chunked text nodes.

        Args:
            url: URL to crawl and process.
            crawl_depth: Depth of crawling (1 = single page, 2 = follow links once).
            extra_metadata: Additional metadata to attach to nodes.

        Returns:
            List of TextNode objects with text chunks and metadata.

        Raises:
            ValueError: If URL is invalid.
            Exception: If crawling fails.
        """
        from datetime import datetime

        from src.document_processing.web_crawler import WebCrawler

        logger.info(f"Processing URL: {url} (depth={crawl_depth})")

        try:
            # Initialize crawler
            crawler = WebCrawler()

            # Crawl URL(s)
            pages = crawler.crawl_url(url, crawl_depth=crawl_depth)

            if not pages:
                logger.warning(f"No content extracted from {url}")
                return []

            all_nodes = []

            # Process each crawled page
            for page_data in pages:
                page_url = page_data["url"]
                content = page_data["content"]
                url_hash = page_data["url_hash"]
                depth = page_data["depth"]

                if not content.strip():
                    logger.warning(f"Empty content from {page_url}")
                    continue

                # Compute content hash (different from URL hash)
                content_hash = hashlib.sha256(content.encode()).hexdigest()

                # Current timestamp for last_crawled
                crawl_timestamp = datetime.now().timestamp()

                # Create base metadata
                metadata = {
                    "filename": f"Web: {page_url[:100]}",  # Truncated URL for display
                    "source_type": "web",
                    "source_url": page_url,
                    "url_hash": url_hash,
                    "content_hash": content_hash,
                    "last_crawled": crawl_timestamp,
                    "crawl_depth": depth,
                    "file_type": ".html",
                }

                # Add extra metadata if provided
                if extra_metadata:
                    metadata.update(extra_metadata)

                # Create Document object for LlamaIndex
                document = Document(text=content, metadata=metadata)

                # Split into chunks
                nodes = self.splitter.get_nodes_from_documents([document])

                all_nodes.extend(nodes)

                logger.info(
                    f"Successfully processed {page_url}: {len(nodes)} chunks created"
                )

            logger.info(
                f"URL processing complete: {len(all_nodes)} total chunks from {len(pages)} pages"
            )

            return all_nodes

        except Exception as e:
            logger.error(f"Failed to process URL {url}: {e}")
            raise

    def process_directory(
        self,
        directory_path: Path,
        source_type: str = "manual",
    ) -> List[TextNode]:
        """
        Process all supported documents in a directory.

        Args:
            directory_path: Path to the directory containing documents.
            source_type: Source of the documents ('manual' or 'email').

        Returns:
            List of all TextNode objects from all documents.
        """
        logger.info(f"Processing directory: {directory_path}")

        all_nodes = []
        supported_extensions = {".pdf", ".docx", ".txt", ".csv", ".xls", ".xlsx"}

        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    nodes = self.process_document(file_path, source_type=source_type)
                    all_nodes.extend(nodes)
                except Exception as e:
                    logger.error(f"Skipping {file_path} due to error: {e}")
                    continue

        logger.info(
            f"Processed {len(all_nodes)} chunks from directory {directory_path}"
        )

        return all_nodes
